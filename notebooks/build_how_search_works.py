"""Generate notebooks/how_search_works.docx from the content below.

Re-run whenever the algorithm in web/lib/search.ts changes. Edit the
constants + section text in this file, run the script, commit both.

    cd api && source .venv/bin/activate
    python ../notebooks/build_how_search_works.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent / "how_search_works.docx"

# ────────────────────────────────────────────────────────────────────────────
# Build
# ────────────────────────────────────────────────────────────────────────────

doc = Document()


def p_with_runs(parts):
    """Add a paragraph composed of (text, {bold/italic}) runs."""
    para = doc.add_paragraph()
    for text, style in parts:
        run = para.add_run(text)
        if style.get("bold"):
            run.bold = True
        if style.get("italic"):
            run.italic = True
    return para


# ─── Title ─────────────────────────────────────────────────────────────────
doc.add_heading("How the Search Algorithm Works", level=0)
sub = doc.add_paragraph()
r = sub.add_run(
    "Companion to algorithm_design.ipynb · Living document — last updated "
    "2026-04-25 · OIT 277 Artist Marketplace project"
)
r.italic = True

doc.add_paragraph(
    "Read top to bottom. No technical background needed. Tables and analogies "
    "do most of the work."
)


# ─── 1. 60-second version ──────────────────────────────────────────────────
doc.add_heading("1. The 60-Second Version", level=1)
doc.add_paragraph(
    "When a buyer types something like “moody lo-fi music for a quiet "
    "vlog under $80,” our algorithm does eight things in order:"
)
for step in [
    "Removes anything that does not match the buyer’s must-haves "
    "(wrong medium, way over budget, not verified when they asked for verified).",
    "Picks the 60 closest matches by meaning — not by keyword. The "
    "system understands that “lo-fi” and “chill” are "
    "related even though they are different words.",
    "Scores each of the 60 on three dimensions: relevance, emerging-artist "
    "boost, and price penalty. (Diversity is handled mechanically in the "
    "next step — it is not a score weight in v1.)",
    "Spreads results across artists by round-robin, so one creator cannot "
    "dominate the top spots.",
    "Returns the top 8 by combined score.",
    "Double-checks that at least 1–2 emerging artists are in the top 8. "
    "If not, swaps in the best emerging artist who passed a relevance bar.",
    "Labels each result with a confidence tag: HIGH, MEDIUM, LOW, or "
    "NO_MATCH.",
    "Surfaces the per-component score breakdown to callers (so the UI can "
    "explain why each result ranked where it did).",
]:
    doc.add_paragraph(step, style="List Number")

p_with_runs([
    ("The point. ", {"bold": True}),
    (
        "The algorithm is not just sorting by relevance. It is executing a "
        "specific strategy: surface emerging artists first, while respecting "
        "what the buyer asked for.",
        {},
    ),
])


# ─── 2. Mission ────────────────────────────────────────────────────────────
doc.add_heading("2. The Mission, in One Sentence", level=1)
p_with_runs([
    (
        "Our mission is to give emerging independent artists exposure. "
        "Relevance is a tool we use to make that happen — not the goal "
        "itself.",
        {"italic": True},
    )
])
doc.add_paragraph(
    "This is a stronger stance than most marketplaces take. Etsy, for "
    "example, optimizes for “what is most likely to be clicked,” "
    "which over time means drop-shippers winning over handmade sellers. We "
    "are deliberately resisting that pattern."
)
doc.add_paragraph(
    "The technical name for our objective is “emerging-artist exposure."
    "” That is what the entire algorithm is engineered around."
)


# ─── 3. Pipeline ───────────────────────────────────────────────────────────
doc.add_heading(
    "3. The Pipeline — What Happens When You Click “Search”",
    level=1,
)
doc.add_paragraph(
    "Imagine the algorithm as a librarian helping you find a book. Here is "
    "what they do, step by step, using the running example:"
)
p_with_runs([
    ("Buyer types: ", {}),
    (
        "“moody lo-fi music for a quiet vlog, budget $65, only verified "
        "artists.”",
        {"italic": True},
    ),
])

# Step 1
doc.add_heading("Step 1 — Listen to the request", level=2)
doc.add_paragraph("The librarian writes down two things separately:")
doc.add_paragraph(
    "The vibe of what they want: “moody lo-fi music for a quiet vlog”",
    style="List Bullet",
)
doc.add_paragraph(
    "The hard rules: must be music; budget $65; verified artists only",
    style="List Bullet",
)
p_with_runs([
    ("Why separate? ", {"bold": True}),
    (
        "Because the librarian treats these differently. The vibe is fuzzy "
        "— “lo-fi” could match many things. The hard rules are "
        "non-negotiable — a buyer who said “music” does not "
        "want a video.",
        {},
    ),
])

# Step 2
doc.add_heading("Step 2 — Apply hard filters", level=2)
doc.add_paragraph(
    "Walking the shelves, the librarian first rules out anything that "
    "breaks the hard rules:"
)
for b in [
    "Books that are not music (illustrations, videos, characters — gone)",
    "Books from non-verified artists (gone)",
    "Books priced above $97.50 (= $65 × 1.5 — anything way over "
    "budget is gone)",
]:
    doc.add_paragraph(b, style="List Bullet")
p_with_runs([
    ("Why $97.50, not $65? ", {"bold": True}),
    (
        "Because the buyer’s “$65 budget” is a preference, not "
        "a hard wall. A truly outstanding $80 work might still be worth "
        "showing. But $200 is too far over. So we draw the line at 1.5× "
        "the budget — far enough to give wiggle room, close enough that "
        "we are still respecting the buyer.",
        {},
    ),
])

# Step 3
doc.add_heading("Step 3 — Find the 60 closest matches", level=2)
doc.add_paragraph(
    "From what is left, the librarian picks the 60 books whose meaning is "
    "closest to “moody lo-fi music for a quiet vlog.”"
)
p_with_runs([
    ("How? ", {"bold": True}),
    (
        "Each book has a meaning fingerprint — a numerical summary of "
        "its mood, style, and topic. The buyer’s query also gets a "
        "fingerprint. The librarian compares fingerprints and picks the 60 "
        "closest.",
        {},
    ),
])
p_with_runs([
    ("Why 60, not 8? ", {"bold": True}),
    (
        "Because the next steps do more careful work, and we want a "
        "high-quality candidate pool with room for the diversity rerank and "
        "quota repair to do their jobs. Like screening 1,000 résumés "
        "down to 60 before doing in-depth interviews on each. (The exact "
        "rule in code: max(K × 6, 60), so for K = 8 we fetch 60.)",
        {},
    ),
])

# Step 4
doc.add_heading("Step 4 — Score each candidate", level=2)
doc.add_paragraph(
    "This is where the strategy lives. For each of the 60 candidates, the "
    "librarian computes a score by combining three things:"
)

# Score table — 3 dimensions
score_tbl = doc.add_table(rows=4, cols=3)
score_tbl.style = "Light Grid Accent 1"
hdr = score_tbl.rows[0].cells
hdr[0].text = "Dimension"
hdr[1].text = "What it asks"
hdr[2].text = "Weight (v1)"
rows = [
    (
        "Relevance (α)",
        "How well does this match what the buyer asked?",
        "1.0",
    ),
    (
        "Health (δ)",
        "Does this give exposure to a new/emerging artist? "
        "(Only counted when relevance ≥ 0.20, so we never boost "
        "low-relevance results just for being emerging.)",
        "0.5",
    ),
    (
        "Price penalty (ε)",
        "How far is this over the buyer’s budget? (subtracted)",
        "0.6",
    ),
]
for i, (a, b, c) in enumerate(rows, start=1):
    cells = score_tbl.rows[i].cells
    cells[0].text = a
    cells[1].text = b
    cells[2].text = c

doc.add_paragraph("The final score is:")
p_with_runs([
    (
        "Score = 1.0 × Relevance + 0.5 × Health "
        "− 0.6 × Price Penalty",
        {"bold": True},
    )
])
doc.add_paragraph("Each dimension is a number between 0 and 1.")

p_with_runs([
    ("A note on diversity. ", {"bold": True}),
    (
        "An earlier draft of this document included a fourth dimension "
        "— Diversity (γ) — that lowered a result’s score "
        "if it was too similar to items already shown. In v1 we replaced "
        "that with a simpler mechanical step: round-robin by artist (Step 5). "
        "Cheaper, more predictable, and γ in the live code is set to 0. "
        "The MMR-style diversity weight is reserved for v2.",
        {},
    ),
])

doc.add_heading(
    "Worked example — three candidates for the running query",
    level=3,
)
doc.add_paragraph(
    "Three candidates the system considered for the running query, scored "
    "with the v1 formula:"
)

ex_tbl = doc.add_table(rows=4, cols=5)
ex_tbl.style = "Light Grid Accent 1"
ex_hdr = ex_tbl.rows[0].cells
ex_hdr[0].text = "Candidate"
ex_hdr[1].text = "Relevance"
ex_hdr[2].text = "Health (emerging?)"
ex_hdr[3].text = "Price"
ex_hdr[4].text = "Final score"
ex_rows = [
    (
        "Aya Tolentino — “Rainy Day Reverie”",
        "0.45",
        "1.00 (emerging, 900 followers)",
        "$40 (no penalty)",
        "0.95",
    ),
    (
        "Noah Patel — “Midnight Loop”",
        "0.57",
        "0.00 (established, 12k followers)",
        "$65 (no penalty)",
        "0.57",
    ),
    (
        "Noah Patel — “Window Light”",
        "0.44",
        "0.00 (established)",
        "$72 (slightly over budget, pen ≈ 0.05)",
        "0.41",
    ),
]
for i, row in enumerate(ex_rows, start=1):
    cells = ex_tbl.rows[i].cells
    for j, val in enumerate(row):
        cells[j].text = val

p_with_runs([
    ("Why Aya wins despite Noah having higher relevance: ", {"bold": True}),
    (
        "She is emerging (900 followers vs Noah’s 12,000) → +0.5 "
        "boost. The math: 1.0 × 0.45 + 0.5 × 1.00 = 0.95.",
        {},
    ),
])
p_with_runs([
    ("Why the third Noah work loses points: ", {"bold": True}),
    (
        "$72 is $7 over the $65 budget. Inside the soft band ($65–$97.50), "
        "the penalty is small (about 0.05) but nonzero. The math: "
        "1.0 × 0.44 + 0.5 × 0 − 0.6 × 0.05 ≈ 0.41.",
        {},
    ),
])

doc.add_heading("The weights are the strategy", level=3)
doc.add_paragraph(
    "The numbers (1.0 / 0.5 / 0.6) are the platform’s strategy in code "
    "form:"
)
for b in [
    "Increase δ → push emerging artists harder.",
    "Increase ε → treat the budget more strictly.",
    "Increase β (currently 0) → reward profitable listings (we are "
    "deliberately keeping this off — see Section 6).",
    "Increase γ (currently 0) → turn embedding-based diversity back "
    "on. We chose round-robin instead for v1.",
]:
    doc.add_paragraph(b, style="List Bullet")
doc.add_paragraph(
    "The product team owns those numbers. Engineers do not change them; "
    "engineers change the formula structure."
)

# Step 5 (NEW)
doc.add_heading("Step 5 — Spread results across artists", level=2)
doc.add_paragraph(
    "Even with strong scores, three works by the same prolific artist can "
    "easily land in the top 8 — not the experience we want. So the "
    "librarian reorders the scored list by round-robin: take the top-scoring "
    "work by Artist A, then Artist B, then Artist C… and keep cycling "
    "until each artist’s queue is empty or the slot count is hit. "
    "Within each artist’s pile the score order is preserved."
)
p_with_runs([
    ("Why mechanical and not weighted? ", {"bold": True}),
    (
        "An earlier design used an embedding-based diversity weight "
        "(γ · MMR). Round-robin is cheaper, more predictable, and "
        "easier to explain. We can promote it back to a score weight in v2 "
        "if we need finer control.",
        {},
    ),
])

# Step 6 (was Step 5)
doc.add_heading("Step 6 — Quota repair (the safety net)", level=2)
doc.add_paragraph(
    "After scoring and round-robin, the librarian looks at the top 8 and "
    "asks: “Do at least 2 emerging artists appear?” (For top-5, "
    "the threshold is 1. The exact rule in code: ceil(K × 0.2), with "
    "a floor of 1.)"
)
for b in [
    "If yes → done.",
    "If no → swap the lowest-ranked non-emerging slot for the "
    "highest-ranked emerging artist outside the top 8 — but only if "
    "that emerging artist passed a relevance bar of 0.40. The repaired row "
    "is flagged with quota_repaired: true so the UI can mark it.",
]:
    doc.add_paragraph(b, style="List Bullet")
p_with_runs([
    ("Why the relevance bar? ", {"bold": True}),
    (
        "Because we don’t want to fill the quota with garbage. "
        "Underfill is better than spam. If no emerging artist passed the "
        "bar, we leave the slot alone.",
        {},
    ),
])
p_with_runs([
    ("Belt and braces. ", {"bold": True}),
    (
        "The score formula already favors emerging artists in Step 4 "
        "(δ = 0.5). The quota is a safety net for the worst case where "
        "the formula alone would not have surfaced any.",
        {},
    ),
])

# Step 7 (was Step 6)
doc.add_heading("Step 7 — Tag confidence levels", level=2)
doc.add_paragraph("Each result gets a label based on its relevance score:")
ct = doc.add_table(rows=5, cols=3)
ct.style = "Light Grid Accent 1"
ct_hdr = ct.rows[0].cells
ct_hdr[0].text = "Tag"
ct_hdr[1].text = "Relevance score"
ct_hdr[2].text = "Meaning"
ct_rows = [
    ("HIGH", "≥ 0.55", "Strong match — the system is confident"),
    ("MEDIUM", "≥ 0.40", "Decent fit"),
    ("LOW", "≥ 0.25", "Thin match — show with caution"),
    (
        "NO_MATCH",
        "below 0.25",
        "No good answer — caller can choose to reject the search and "
        "tell the user",
    ),
]
for i, row in enumerate(ct_rows, start=1):
    for j, val in enumerate(row):
        ct.rows[i].cells[j].text = val

doc.add_paragraph(
    "The frontend uses these tags to decide whether to label a result "
    "“Recommended” or “Best we could find” — or to "
    "show a “no match” message instead."
)

# Step 8 (replaces old Step 7 logging — softened to per-result breakdown)
doc.add_heading("Step 8 — Surface the score breakdown", level=2)
doc.add_paragraph(
    "Every returned row carries its component scores back to the caller: "
    "the raw relevance (rel), health bonus (health), price penalty "
    "(price_pen), and the final composite (score). The frontend can "
    "render any of these (e.g., “92% match” uses rel)."
)
p_with_runs([
    ("Why this matters. ", {"bold": True}),
    (
        "If a teammate or a buyer ever asks “why did you rank these "
        "this way?” the breakdown is right there in the response. "
        "Structured request-level audit logging — capturing the full "
        "candidate pool, filters, and swaps for replay — is on the "
        "deferred list (Section 6).",
        {},
    ),
])


# ─── 4. Strategic Decisions ────────────────────────────────────────────────
doc.add_heading("4. The Strategic Decisions Behind the Numbers", level=1)
doc.add_paragraph(
    "The algorithm above is one specific implementation of a broader "
    "strategy. Here are the calls we made and why:"
)
sd = doc.add_table(rows=10, cols=3)
sd.style = "Light Grid Accent 1"
sd_hdr = sd.rows[0].cells
sd_hdr[0].text = "Decision"
sd_hdr[1].text = "Choice"
sd_hdr[2].text = "Why"
sd_rows = [
    ("Lead metric", "Emerging-artist exposure", "The platform mission — not a bolt-on"),
    (
        "Medium and verified-only",
        "HARD requirements (must-pass)",
        "A buyer who says “music” does not want video; "
        "“verified” is a yes/no question",
    ),
    (
        "Price",
        "SOFT within 1.5× budget; HARD beyond",
        "An outstanding match at $80 against a $65 budget should still "
        "surface; $200 should not",
    ),
    (
        "“Emerging” artist threshold",
        "reach_score below 1,500 followers",
        "Matches the bucket already used in our pricing logic — single "
        "source of truth",
    ),
    (
        "Quota size",
        "1 of top-5; 2 of top-10",
        "Modest safety net; the searcher’s needs lead",
    ),
    (
        "Quota relevance floor",
        "Relevance ≥ 0.40 to fill a quota slot",
        "Don’t fill the quota with garbage — underfill beats spam",
    ),
    (
        "δ weight (emerging boost)",
        "0.5 (moderate)",
        "Real boost without overriding the searcher",
    ),
    (
        "γ weight (diversity)",
        "0 — round-robin by artist instead",
        "Cheaper and easier to explain than embedding-based MMR; can promote "
        "to a score weight in v2",
    ),
    (
        "β weight (revenue)",
        "0 (off in v1)",
        "Etsy failure mode — turn on only after we have data",
    ),
]
for i, (a, b, c) in enumerate(sd_rows, start=1):
    cells = sd.rows[i].cells
    cells[0].text = a
    cells[1].text = b
    cells[2].text = c


# ─── 5. How we know it is working ──────────────────────────────────────────
doc.add_heading("5. How We Know It Is Working", level=1)
doc.add_paragraph(
    "We measure quality on a small set of hand-labeled “golden” "
    "queries — queries where we manually wrote down which results "
    "SHOULD appear. The most recent measurements (against 6 golden queries "
    "on the seed catalog of 148 works) are reproduced below."
)
p_with_runs([
    ("Note: ", {"bold": True}),
    (
        "the numbers below were measured before γ went to 0 in "
        "production. They will move when re-measured against the live v1 "
        "ranker. Treat as directional until the next eval pass appends a "
        "fresh row to the change log.",
        {"italic": True},
    ),
])

ev = doc.add_table(rows=7, cols=3)
ev.style = "Light Grid Accent 1"
ev_hdr = ev.rows[0].cells
ev_hdr[0].text = "Metric"
ev_hdr[1].text = "What it means"
ev_hdr[2].text = "Latest result"
ev_rows = [
    ("Recall@8", "Of the works we said are relevant, what % made it into the top 8?", "1.000 (100%)"),
    ("Precision@8", "Of the top 8, what % were genuinely relevant?", "0.375 (≈ 3 of 8)"),
    ("MRR", "How fast does the first useful result appear? (1.0 = always #1)", "0.694 (≈ position 1.4)"),
    ("NDCG@8", "Are the top spots filled with the strongest matches?", "0.724"),
    ("Emerging share", "What % of all results shown were emerging artists?", "30.4%"),
    ("Price fit", "What % of results respected the buyer’s budget?", "100%"),
]
for i, (a, b, c) in enumerate(ev_rows, start=1):
    cells = ev.rows[i].cells
    cells[0].text = a
    cells[1].text = b
    cells[2].text = c

p_with_runs([
    ("Why three layers of evaluation? ", {"bold": True}),
    (
        "The numbers above are “Layer 1” (math metrics on a "
        "labeled test set). We also run “Layer 2” (marketplace-"
        "level metrics like emerging share, distinct artists shown, "
        "concentration) and we plan to add “Layer 3” "
        "(rubric-based judgment by a human or AI) once the team grows.",
        {},
    ),
])


# ─── 6. What's NOT in the algorithm yet ────────────────────────────────────
doc.add_heading("6. What Is NOT in the Algorithm Yet (and Why)", level=1)
doc.add_paragraph(
    "Four things we deliberately deferred. Each will turn on later, under "
    "specific conditions."
)

doc.add_heading("β — Revenue weight", level=3)
doc.add_paragraph(
    "Right now β = 0, meaning we ignore profitability when ranking. "
    "We are keeping it off because turning revenue on prematurely creates "
    "the Etsy failure mode (drop-shippers outrank handmade sellers)."
)
p_with_runs([
    ("Turn on when: ", {"bold": True}),
    (
        "we have real conversion data AND enough artist density that the "
        "long tail can afford some revenue pressure.",
        {},
    ),
])

doc.add_heading("γ — Embedding-based diversity weight", level=3)
doc.add_paragraph(
    "Right now γ = 0; diversity is handled mechanically by round-robin "
    "(Step 5). An MMR-style γ weight would let two near-identical "
    "works penalize each other within the score itself — finer "
    "control, but more expensive and harder to explain."
)
p_with_runs([
    ("Turn on when: ", {"bold": True}),
    (
        "round-robin produces visibly clustered results in real searches, "
        "or when stylistic diversity (within an artist’s catalog) "
        "becomes the bottleneck.",
        {},
    ),
])

doc.add_heading("Cold-start exploration", level=3)
doc.add_paragraph(
    "Today, an emerging artist with no track record can still surface via "
    "the quota — but only if they pass the relevance bar. We do not "
    "deliberately surface random new works to learn what works."
)
p_with_runs([
    ("Turn on when: ", {"bold": True}),
    (
        "the catalog grows past ~1,000 works, where the quota alone stops "
        "being enough to surface the long tail.",
        {},
    ),
])

doc.add_heading("Structured request-level audit logging", level=3)
doc.add_paragraph(
    "Every search currently returns its per-result score breakdown to the "
    "caller (Step 8), but we do not persist a full audit row capturing the "
    "candidate pool, hard-filter rejections, swap decisions, and "
    "post-quota state for replay. That is a separate piece of plumbing."
)
p_with_runs([
    ("Turn on when: ", {"bold": True}),
    (
        "we need to debug specific buyer complaints (“why did this "
        "show?”) at scale, or when an algorithm change requires "
        "comparing rankings before-and-after on real traffic.",
        {},
    ),
])

doc.add_heading("AI-judge evaluation", level=3)
doc.add_paragraph(
    "Today, our test set is 6 manually-labeled queries. As the team grows, "
    "manually inspecting every algorithm change becomes unsustainable."
)
p_with_runs([
    ("Turn on when: ", {"bold": True}),
    (
        "we are making more than 1–2 algorithm changes per week, and "
        "the test set needs to grow past what humans can label.",
        {},
    ),
])


# ─── 7. Glossary ───────────────────────────────────────────────────────────
doc.add_heading("7. Glossary", level=1)
gl = doc.add_table(rows=12, cols=2)
gl.style = "Light Grid Accent 1"
gl_hdr = gl.rows[0].cells
gl_hdr[0].text = "Term"
gl_hdr[1].text = "Plain English"
gl_rows = [
    ("Catalog", "The full set of artworks listed on the platform"),
    ("Query", "The buyer’s text input (e.g., “moody lo-fi music”)"),
    ("Hard filter", "A rule that absolutely excludes something (e.g., wrong medium)"),
    ("Soft constraint", "A penalty that lowers a score but does not exclude an item"),
    ("Top-K", "The K results we show the buyer (in v1, K = 8)"),
    ("Emerging artist", "An artist with reach_score below 1,500 followers"),
    (
        "Reach score",
        "Total followers across an artist’s linked socials "
        "(Instagram + TikTok + website)",
    ),
    ("Quota", "A guarantee that at least N emerging artists appear in the top-K"),
    (
        "Relevance floor",
        "A minimum relevance score below which a result will not be shown / "
        "will not fill a quota slot",
    ),
    (
        "Score breakdown",
        "Per-result component scores (rel, health, price_pen) returned with "
        "every search",
    ),
    (
        "Confidence tier",
        "A label (HIGH / MEDIUM / LOW / NO_MATCH) telling the UI how strong "
        "the match is",
    ),
]
for i, (a, b) in enumerate(gl_rows, start=1):
    gl.rows[i].cells[0].text = a
    gl.rows[i].cells[1].text = b


# ─── 8. Change log ─────────────────────────────────────────────────────────
doc.add_heading("8. Change Log", level=1)
doc.add_paragraph("Every algorithm change appends a row here. New items go at the top.")
cl = doc.add_table(rows=4, cols=3)
cl.style = "Light Grid Accent 1"
cl_hdr = cl.rows[0].cells
cl_hdr[0].text = "Date"
cl_hdr[1].text = "What changed"
cl_hdr[2].text = "Why"
cl_rows = [
    (
        "2026-04-25",
        "Doc aligned with the shipped v1 ranker. Removed γ (diversity) "
        "from the score formula — it is 0 in production, since "
        "round-robin by artist handles diversity mechanically; γ is "
        "now described in Section 6 as a deferred weight. Pipeline split "
        "from 7 steps into 8 (round-robin gets its own step). Worked-example "
        "scores recomputed with the v1 formula. Candidate pool size updated "
        "to 60 (was described as 30; code uses max(K × 6, 60)). "
        "Replaced the old “Log everything” step with “Surface "
        "the score breakdown” — honest about what the code does. "
        "Structured request-level audit logging moved to Section 6 "
        "(deferred).",
        "Living-doc rule: doc tracks the code. Caught during a sync pass "
        "after the v1 deploy.",
    ),
    (
        "2026-04-25",
        "Algorithm v1 deployed to production (web/lib/search.ts). Live "
        "search endpoints /api/search and /api/search/artists now apply the "
        "full pipeline: hard filters (medium, verified-only, price ceiling "
        "at 1.5x budget) → weighted score (α·rel + "
        "δ·health − ε·price_pen, δ=0.5, "
        "ε=0.6) → round-robin artist diversity → quota "
        "repair (1/5 or 2/10 emerging with rel ≥ 0.40) → "
        "confidence tagging. Frontend chips parsed into structured filters "
        "(medium/budget_cents/verified_only) instead of free-text "
        "concatenation.",
        "Prove the algorithm in the live UI. γ (MMR-style diversity) "
        "stays 0 in production v1 because we don’t fetch embeddings "
        "client-side; round-robin by artist_id is the diversity mechanism. "
        "β (revenue) still 0. SQL migration for filter-aware "
        "search_works RPC is deferred (filters applied in TypeScript after "
        "fetching candidate metadata).",
    ),
    (
        "2026-04-25",
        "Initial version of this document. Algorithm v1: weights α=1.0, "
        "γ=0.4, δ=0.5, ε=0.6; soft price band (1.5× "
        "budget ceiling, hard reject beyond); quota 1/5 or 2/10 with "
        "relevance floor 0.40; emerging = reach_score < 1,500.",
        "Strategic stance locked via group discussion: “emerging-first "
        "with searcher’s needs respected.” β stays off.",
    ),
]
for i, (a, b, c) in enumerate(cl_rows, start=1):
    cells = cl.rows[i].cells
    cells[0].text = a
    cells[1].text = b
    cells[2].text = c


# ─── Closer ────────────────────────────────────────────────────────────────
doc.add_heading("A note on this document", level=1)
doc.add_paragraph(
    "This is a living document. Every time the algorithm in "
    "algorithm_design.ipynb or web/lib/search.ts changes, the relevant "
    "section here will be updated and a new row will be added to the "
    "change log. If you ever find a discrepancy between what the code does "
    "and what this document says, that is a bug — flag it and we will "
    "fix it."
)
doc.add_paragraph(
    "No part of this document requires technical background to read. If "
    "anything in here is unclear, the right move is to push back on the "
    "writing, not to assume you are missing something."
)


# ─── Save ──────────────────────────────────────────────────────────────────
doc.save(str(OUT))
print(f"wrote {OUT.name}: {OUT.stat().st_size // 1024} KB")
